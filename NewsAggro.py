# News Aggrigator 
# Input comma separated list, search those terms on Multiple news sources
# Output word doc of 10 articles for each thing formatted as follows:
# Keyword:
# Title
# First 25(ish) words ...
# Link to article

# Powered in part by NewsAPI.org

import docx
import requests
import easygui
import json
import feedparser
import time
import queue    # PyInstaller dependancy

exclude = ['breitbart-news']    # Exclude list from NewsAPI

                                # Include list of rss feeds formatted [ (name, link) ]
                                # Note, some rss feeds aren't formatted very well with feedparser. 
                                # May have to manually make new functions to parse them
rssFeeds = [('cnn-tech', 'http://rss.cnn.com/rss/cnn_tech.rss'), ('fox-news', 'http://feeds.foxnews.com/foxnews/latest?format=xml'), 
            ('legal-tech-news', 'http://feeds.feedblitz.com/lawtechnologynews'), ('dhs', 'https://www.dhs.gov/news/rss.xml'), 
            ('ics-cert', 'https://ics-cert.us-cert.gov/xml/rss.xml'), ('naked-security', 'https://nakedsecurity.sophos.com/feed/'),
            ('threat-post', 'https://threatpost.com/feed/'), ('CVEs', 'https://nvd.nist.gov/download/nvd-rss-analyzed.xml'),
            ('acm-tech-news', 'http://rss.acm.org/technews/TechNews.xml'), ('breaking-defense', 'http://feeds.feedburner.com/BreakingDefense?format=xml'),
            ('defense.gov', 'https://www.defense.gov/DesktopModules/ArticleCS/RSS.ashx?max=15&ContentType=9&Site=727'),
            ('security-magazine-physical', 'http://www.securitymagazine.com/rss/topic/2228'), ('security-magazine-cyber', 'http://www.securitymagazine.com/rss/topic/2788'),
            ('professional-security', 'http://www.professionalsecurity.co.uk/rss-feed/')]

API = '' # Insert NewsAPI key here
Intel471_API = 'Basic ' # Insert Base 64 encoded User and API key here

def clean_breaking_defense(f):
    oldDesc = f['description']
    newDesc = ''

    #always the second <p> tag in the description
    splitDesc = oldDesc.split('</p>')
    newDesc = splitDesc[1][3:]

    f['description'] = newDesc

    return f

def clean_naked_sec(f):
    oldDesc = f['description']
    newDesc = ''
    
    for char in oldDesc:
        if char != '<':
            newDesc += char
        else:
            break

    f['description'] = newDesc
    return f

def clean_legal_tech_news(f):
    charNum = 0
    for letter in f['description']:
        if letter == '<':
            f['description'] = f['description'][:charNum]
            break
        charNum += 1

    return f

def clean_dhs(f, cutoff=300):
    oldDesc = f['description'].split('<p>')
    newParagraph = ''
    newDesc = ''

    for paragraph in oldDesc:
        if ((len(paragraph) > 75) and (paragraph.startswith('<div') == False) and 
        ('FOR IMMEDIATE RELEASE' not in paragraph.upper())):
            newParagraph = paragraph
            break
    
    for char in newParagraph:
        if char != '<':
            newDesc += char
        else:
            break
    
    if len(newDesc) > cutoff:
        newDesc = newDesc[0:cutoff] + '...'
    f['description'] = newDesc
    return f

def get_rss_feed(url, publication, numArticles=15):
    d = feedparser.parse(url)
    entryList = []

    for entry in range(numArticles):
        try:
            retVal = {  'publication': publication,
                        'title': d.entries[entry].title,
                        'description': d.entries[entry].description,
                        'url': d.entries[entry].link
                        #'pubDate': d.entries[entry].pubDate
                        }
            if publication == 'legal-tech-news':
                retVal = clean_legal_tech_news(retVal)
            elif publication == 'dhs':
                retVal = clean_dhs(retVal)
            elif publication == 'naked-security':
                retVal = clean_naked_sec(retVal)
            elif publication == 'breaking-defense':
                retVal = clean_breaking_defense(retVal)

            entryList.append(retVal)
        
        except IndexError:      # Ran out of RSS entries to read in
            return entryList

    return entryList

def get_intel471_json(apiKey, sort='latest', count='25'):
    url = 'https://api.intel471.com/v1/reports'
    params = {  'sort': sort,
                'count': count,
                'prettyPrint': True}
    
    header = {  'Authorization': apiKey }

    r = requests.get(url, params=params, headers=header).json()
    return r['reports']

def get_json_response(  apiKey, source='google-news', sortBy='top', category='politics',
                        language='en', country='us', request='articles'):
    url = 'https://newsapi.org/v1/' + request
    
    if request == 'articles':
        params = { 'source': source,
                    'apiKey': apiKey,
                    'sortBy': sortBy}
    
    elif request == 'sources':
        params = {  'category': category,
                    'language': language} 
                 #, 'country': country} Usually this is overly limiting

    r = requests.get(url, params=params).json()
    
    try:
        if r['code'] == 'sourceUnavailableSortedBy':    # Sorts by top by default
            r = get_json_response(apiKey, source=source, sortBy='latest')
    except:
        # No error
        pass

    return r

def get_sources(apiKey):
    source_list = []
    categories = ['general', 'politics', 'technology', 'business', 'science-and-nature']
                 # Categories left out: entertainment, gaming, music, sport

    for c in categories:
        response = get_json_response(apiKey, category=c, request='sources')
        for source in response['sources']:
            if source['id'] not in exclude:
                source_list.append(source['id'])

    return source_list

def parse_Intel471_response(api, search_terms):
    hits = {}
    print("Searching  Intel-471 (This may take a while)")
    dict_list = get_intel471_json(api, count=15)

    for article in dict_list:
        for terms in search_terms:
            dict_key = ''
            for term in terms:
                dict_key += term
            
            hit = True
            for term in terms:
                if term not in article['subject'].upper():
                    hit = False
            
            if hit == True:
                if dict_key not in hits:
                    hits[dict_key] = [{
                        'title': article['subject'],
                        'url': article['portalReportUrl'],
                        'publication': 'Intel-471'
                    }]

                else:
                    hits[dict_key].append({
                        'title': article['subject'],
                        'url': article['portalReportUrl'],
                        'publication': 'Intel-471'
                    })

    return hits

def parse_json_sources(apiKey, source_list, search_terms):
    raw_responses = []
    hits = {}

    for source in source_list:
        raw_responses.append(get_json_response(apiKey, source=source))
        print("Searching  " + source)

    source_index = 0     # Keeps track of which source we're on
    for articles in raw_responses:
        source = source_list[source_index]
        for article in articles['articles']:
            try:
                desc = article['description'].upper()
            except KeyError:
                desc = ''
            except AttributeError:
                desc = ''

            try: 
                title = article['title'].upper()
            except KeyError:
                title = ''
            except AttributeError:
                title = ''

            for terms in search_terms:
                dict_key = ''
                
                for term in terms:
                    dict_key += term + ' '
                dict_key = dict_key[:-1]

                append_list = True
                for term in terms:
                    if (term not in title) and (term not in desc):
                        append_list = False

                if append_list:
                    article['publication'] = source
                    if term not in hits:
                        hits[dict_key] = [article]
                    else:
                        hits[dict_key].append(article)
        source_index += 1

    return hits

def parse_rss_sources(feeds, terms):
    rawResponses = []
    hits = {}
    for source in feeds:
        print("Searching  " + source[0])
        rawResponses += get_rss_feed(source[1], source[0])

    for article in rawResponses:
        dict_key = ''
        for searchterms in terms:
            for term in searchterms:
                dict_key += term + ' '
            dict_key = dict_key[:-1]
            
            hit = True
            for term in searchterms:
                if (term not in article['title'].upper()) and (term not in article['description'].upper()):
                    hit = False

            if hit == True:
                if dict_key not in hits:
                    hits[dict_key] = [article]
                else:
                    hits[dict_key].append(article)

    return hits

def format_hits(hits):
    # Writes to word document
    
    d = docx.Document()
    for k, v in hits.items():
        p = d.add_paragraph()
        p.add_run(k + '\n').bold = True

        for article in v:
            p.add_run('Source: ').bold = True
            p.add_run(article['publication'] + '\n')
            
            p.add_run("Title").bold = True
            p.add_run(": " + article['title'] + '\n')

            p.add_run('Description').bold = True
            try:
                p.add_run(': ' + article['description'] + '\n')
            except:
                p.add_run(": No description available\n")

            p.add_run("Link").bold = True
            p.add_run(": " + article['url'] + '\n\n')

    return(d)

def combine_dicts(d1, d2):
    for k1, v1 in d1.items():   # Assimilate all similar elements
        if k1 in d2:
            d1[k1] += d2[k1]
    
    for k2, v2 in d2.items():   # Add all disimilar elements to d1 as well
        if k2 not in d1:
            d1[k2] = d2[k2]
    
    return d1

def save(doc):
    filename = easygui.filesavebox('Save as')
    
    while filename == None:
        quit = easygui.ynbox('Are you sure you want to quit without saving? ')
        if quit == True:
            print('Goodbye!')
            return None
        else:
            filename = easygui.filesavebox('Save as')

    if filename[-5:] != '.docx':
        filename += '.docx'

    try:
        doc.save(filename)
    except:
        easygui.msgbox("Filename or save location not allowed.")
        save(doc)

def getTermsFromFile():     # Splits queries by newline, treating spaces between words as AND Operators
    fname = easygui.fileopenbox("Select file with search terms separated by newlines")
    with open(fname, 'r') as f:
        raw_terms = f.read()

    split_terms = raw_terms.split('\n')
    logic_terms = []

    for terms in split_terms:
        line_terms = []
        for term in terms.split(' '):
            if term != '':
                line_terms.append(term.upper())
        if line_terms not in logic_terms and line_terms != []:
            logic_terms.append(line_terms)

    print(logic_terms)
    return logic_terms

def getTermsFromUser():
    raw_terms = easygui.enterbox("Input each term seperated by commas: ")
    less_raw_terms = raw_terms.split(',')
    search_terms_no_logic = []
    search_terms = []

    for t in less_raw_terms:
        if t != '':
            search_terms_no_logic.append(t.upper())

    for t in search_terms_no_logic:
        split_terms = t.split(' ')
        for term_index in range(len(split_terms)):
            try:
                if split_terms[term_index] == '':
                    split_terms.pop(term_index)
            except (IndexError):
                pass

        if split_terms not in search_terms:
            search_terms.append(split_terms)
    print(search_terms)

    return search_terms

def main():
    global API
    rssTime = time.strftime("%a, %d %b %Y") # Currently not implimented, formats date like xml <pubDate> tag
                                            # ie Mon, 01 Mar 2017
    termsFromFile = easygui.ynbox("Are search terms coming from a file?")
    if termsFromFile == True:
        search_terms = getTermsFromFile()
    else:
        search_terms = getTermsFromUser()
    

    source_list = get_sources(API)
    search_results_json = parse_json_sources(API, source_list, search_terms)
    search_results_rss = parse_rss_sources(rssFeeds, search_terms)
    search_results_intel471 = parse_Intel471_response(Intel471_API, search_terms)

    search_results = combine_dicts(search_results_json, search_results_rss)
    search_results = combine_dicts(search_results, search_results_intel471)
    
    out_doc = format_hits(search_results)
    save(out_doc)

main()

# TODO:    Look for terms as words with spaces before or after
